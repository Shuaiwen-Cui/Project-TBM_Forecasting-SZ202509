# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

MANUSCRIPT = Path(r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx")
RED = RGBColor(255, 0, 0)

refs = [
    "[25]\tZHANG Y X, REN X H, ZHANG J X, ZHANG Y Z, MA Z C. A novel workflow including denoising and hybrid deep learning model for shield tunneling construction parameter prediction[J]. Engineering Applications of Artificial Intelligence, 2024, 133: 108103. DOI: 10.1016/j.engappai.2024.108103.",
    "[26]\tDAI L F, CHEN W M, XIAO M Q, SUN W H, WANG Z Z. Prediction of super-large diameter shield attitude based on LSTM-Transformer[J]. Scientific Reports, 2025, 15: 15725. DOI: 10.1038/s41598-025-98428-8.",
    "[27]\tYAO C R, KONG X X, TANG L, LING X Z, TANG W C. A multivariate time series prediction model for TBM excavation parameters using a Convolution-GRU-Attention neural network[J]. Applied Sciences, 2026, 16(6): 2964. DOI: 10.3390/app16062964.",
]

doc = Document(MANUSCRIPT)
existing = "\n".join(p.text for p in doc.paragraphs)
for ref in refs:
    if ref[:12] not in existing:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.color.rgb = RED
doc.save(MANUSCRIPT)
print("Appended missing references.")
