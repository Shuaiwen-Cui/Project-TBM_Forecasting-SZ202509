# -*- coding: utf-8 -*-
"""Remove response-style wording from manuscript parameter paragraph."""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

MANUSCRIPT = Path(r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx")
RED = RGBColor(255, 0, 0)


def set_para(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].font.color.rgb = RED
    else:
        run = paragraph.add_run(text)
        run.font.color.rgb = RED


doc = Document(MANUSCRIPT)
new_text = (
    "模型结构与超参数经代码核对如下：输入 $X\\in\\mathbb{R}^{T\\times F}$ 经线性嵌入（隐维度 $d=64$）与正弦位置编码后，"
    "送入 2 层 Transformer 编码器（4 头注意力，前馈维度 256，Dropout 0.1，Post-LN），得到最后一时刻隐向量 $h_T$，"
    "再经线性层映射为 $\\hat{Y}\\in\\mathbb{R}^{H\\times F}$。训练采用 Adam 优化器、学习率 1×10⁻³、批大小 64、20 个 epoch。"
    "上述参数与代码中的 `TRANSFORMER_CONFIG` 和 `TransformerForecaster` 实现保持一致。"
)
for p in doc.paragraphs:
    if p.text.startswith("模型结构与超参数经代码核对如下"):
        set_para(p, new_text)
        break
else:
    raise RuntimeError("parameter paragraph not found")
doc.save(MANUSCRIPT)
print("Polished parameter paragraph.")
