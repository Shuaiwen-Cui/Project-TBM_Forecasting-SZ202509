# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

path = Path(r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx")
doc = Document(str(path))

red_runs = 0
black_runs = 0
red_chars = 0
black_chars = 0
red_samples = []

for p in doc.paragraphs:
    for r in p.runs:
        if not r.text.strip():
            continue
        c = r.font.color.rgb
        if c == RGBColor(0xFF, 0x00, 0x00):
            red_runs += 1
            red_chars += len(r.text)
            if len(red_samples) < 15:
                red_samples.append(r.text[:120])
        else:
            black_runs += 1
            black_chars += len(r.text)

print(f"red runs: {red_runs}, chars: {red_chars}")
print(f"black runs: {black_runs}, chars: {black_chars}")
print("--- red samples ---")
for s in red_samples:
    print(s)

# key checks
checks = [
    "基于多头自注意力机制",
    "d=64",
    "如图 2 所示，四类精度",
    "刀盘转速（右上）",
    "离线回放",
    "ZHANG Y X",
    "d=4",
    "如图 3 所示，四类精度",
    "土舱土压（右上）",
]
full = "\n".join(p.text for p in doc.paragraphs)
for c in checks:
    print(f"{c}: {'YES' if c in full else 'NO'}")
