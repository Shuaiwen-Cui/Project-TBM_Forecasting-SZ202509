# -*- coding: utf-8 -*-
"""Deepen manuscript revisions and create point-to-point response."""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

BASE = Path(r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission")
WORK = BASE / "_revision_work"
MANUSCRIPT = BASE / "基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx"
RESPONSE = BASE / "gl2026-1222审稿意见-答复.docx"
MD = BASE / "审稿意见修改说明-gl2026-1222.md"
RIS = WORK / "new_references_for_zotero_import.ris"

RED = RGBColor(255, 0, 0)


def red(run):
    run.font.color.rgb = RED


def set_para(paragraph, text, mark=True):
    if paragraph.text == text:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        if mark:
            red(paragraph.runs[0])
    else:
        run = paragraph.add_run(text)
        if mark:
            red(run)
    return True


def find_para(doc, marker):
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            return paragraph
    raise RuntimeError(f"marker not found: {marker}")


def insert_after(paragraph, text, mark=True):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    run = p.add_run(text)
    if mark:
        red(run)
    return p


def clear_doc(doc):
    body = doc._body._element
    for child in list(body):
        body.remove(child)


def add_para(doc, text="", bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def update_manuscript():
    doc = Document(MANUSCRIPT)

    intro = (
        "当前时序预测前沿集中在数据驱动与深度学习：工业物联网积累了大量多变量时序数据，长程依赖建模、多时间尺度预测及 T–H 配置的联合优化成为热点[6–9]。"
        "既有方法大致分为：（1）机理–数据混合模型：可解释性强、工况稳定时有效，但对复杂耦合与长程依赖依赖机理完备性。"
        "（2）统计模型[10–13]（ARIMA、卡尔曼滤波等）：形式简洁、计算量小，线性或高斯假设对多变量非线性与长程依赖适应性有限。"
        "（3）浅层机器学习[14,15]（SVM[16]、随机森林等）：能捕捉非线性，短时预测有效，对多变量联合表示与长序列依赖能力不足，T、H 选型多依赖经验[17–19]。"
        "（4）循环网络[20]（LSTM[21]/GRU）：通过隐状态刻画时序依赖，但训练串行、长序列上梯度易衰减，且多数工作仅针对单一或少量 (T, H) 评估[22,23]。"
        "（5）Transformer 及其变体：自注意力在 $O(T^2)$ 复杂度下直接建模任意位置对，无需递归即可实现长程依赖[24]；多头设计与编码器并行计算亦源于 Transformer 架构[24]。"
        "近年来，盾构/TBM 领域已有研究开始引入注意力混合模型、LSTM-Transformer 以及 CNN/GRU/Attention 等结构，用于盾构姿态、贯入度、推力或刀盘扭矩等参数预测[25–27]。"
        "这些工作说明注意力机制在施工参数时序建模中具有应用潜力，但多聚焦单一目标、单一预测跨度或少量模型对比；针对同一盾构区间多变量关键参数，在统一直接多步架构下系统扫描输入长度 T 与预测长度 H，并同时给出实时计算开销估计的研究仍不足。"
        "因此，本文侧重多变量联合、直接多步输出与 T–H 联合扫描，并将精度衰减、计算成本和实时滚动调用需求放在同一框架中讨论。"
    )
    set_para(find_para(doc, "当前时序预测前沿集中在数据驱动与深度学习"), intro)

    model_para = (
        "本文采用仅含编码器的注意力模型进行多变量多步预测：输入经线性嵌入与位置编码后，经多层多头自注意力编码器得到最后一时刻隐向量，"
        "再经单层线性映射直接输出未来 H 步（直接多步）。该设计与 Vaswani 等提出的 Transformer 编码器思想一致[24]，但本文保留的是面向盾构多变量滚动预测的核心结构，"
        "即“历史窗口编码—当前状态压缩—未来多步直接输出”。由于预测过程不依赖递归生成，单次前向传播延迟确定，适合在连续采样场景下作为滚动调用的核心预测器；"
        "同时，所有未来步共享同一上下文向量，这也使其在大 H 下存在可解释的结构瓶颈。"
        "因此，本文不是单纯介绍自注意力原理，而是利用这一结构在统一实验中考察 T、H 对精度、时延和存储开销的共同影响。"
    )
    set_para(find_para(doc, "本文采用仅含编码器的注意力模型进行多变量多步预测"), model_para)

    structure_para = (
        "模型结构与超参数经代码核对如下：输入 $X\\in\\mathbb{R}^{T\\times F}$ 经线性嵌入（隐维度 $d=64$）与正弦位置编码后，"
        "送入 2 层 Transformer 编码器（4 头注意力，前馈维度 256，Dropout 0.1，Post-LN），得到最后一时刻隐向量 $h_T$，"
        "再经线性层映射为 $\\hat{Y}\\in\\mathbb{R}^{H\\times F}$。训练采用 Adam 优化器、学习率 1×10⁻³、批大小 64、20 个 epoch。"
        "上述参数与代码 `TRANSFORMER_CONFIG` 和 `TransformerForecaster` 实现一致，纠正了原文中将隐维度误写为 $d=4$ 的问题。"
    )
    set_para(find_para(doc, "模型结构与超参数：输入"), structure_para)

    compute_para = find_para(doc, "计算性能评估在同一计算环境下进行")
    realtime_para = (
        "进一步面向实时在线调用估计计算开销：20 组配置的模型参数量为 104160～850880，按 float32 存储约 0.40～3.25 MB；"
        "单样本前向推理时间约 0.47～3.68 ms，即使取最大配置也远小于 5 s 采样周期（约占 0.074%）。"
        "在线滚动时需要缓存的输入窗口规模为 $T\\times F$，在最大 $T=360,F=32$ 时约 11520 个浮点数，float32 约 45 KB；"
        "若输出 $H=360,F=32$，预测结果缓存同样约 45 KB。由此可见，模型权重存储与单次输入/输出缓存均处于 MB/KB 量级，"
        "实时性主要受大 T、大 H 下批量推理和系统集成开销影响；对于本文建议的较短 H 滚动调用，计算速度和存储开销具备工程实时应用的可行性。"
    )
    insert_after(compute_para, realtime_para)

    refs = {
        "[25]": "ZHANG Y X, REN X H, ZHANG J X, ZHANG Y Z, MA Z C. A novel workflow including denoising and hybrid deep learning model for shield tunneling construction parameter prediction[J]. Engineering Applications of Artificial Intelligence, 2024, 133: 108103. DOI: 10.1016/j.engappai.2024.108103.",
        "[26]": "DAI L F, CHEN W M, XIAO M Q, SUN W H, WANG Z Z. Prediction of super-large diameter shield attitude based on LSTM-Transformer[J]. Scientific Reports, 2025, 15: 15725. DOI: 10.1038/s41598-025-98428-8.",
        "[27]": "YAO C R, KONG X X, TANG L, LING X Z, TANG W C. A multivariate time series prediction model for TBM excavation parameters using a Convolution-GRU-Attention neural network[J]. Applied Sciences, 2026, 16(6): 2964. DOI: 10.3390/app16062964.",
    }
    last_ref = find_para(doc, "[24]")
    # Remove any old bad [25] if still present, then append current refs.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("[25]"):
            set_para(paragraph, refs["[25]"])
            break
    else:
        insert_after(last_ref, refs["[25]"])
    p25 = find_para(doc, refs["[25]"][:30])
    insert_after(p25, refs["[26]"])
    p26 = find_para(doc, refs["[26]"][:30])
    insert_after(p26, refs["[27]"])

    doc.save(MANUSCRIPT)


def update_response_doc():
    doc = Document(RESPONSE)
    clear_doc(doc)
    add_para(doc, "gl2026-1222 审稿意见逐条答复", bold=True)
    add_para(doc, "稿件题目：面向实时在线应用的盾构掘进关键参数多步预测研究")
    add_para(doc, "作者：李奇涛，姜红峰，崔帅文，徐金峰")
    add_para(doc, "")
    add_para(doc, "尊敬的编辑和审稿专家：")
    add_para(doc, "感谢您对本文的认真审阅和建设性意见。我们已根据意见对题目、摘要、引言、方法、计算性能分析、图题图号、模型参数和参考文献进行了逐条修改。修改稿中新增或改写内容已用红色字体标出。以下为逐条答复。")
    add_para(doc, "")

    items = [
        (
            "意见1：文章题目、摘要及正文多处强调“实时预测、在线部署、滚动更新、跨区间迁移”等内容，但实际开展的是基于单一区间数据的离线直接多步预测实验，未进行实时数据流测试、在线滚动预测验证以及跨区间迁移实验，相关表述与实际实验内容之间不匹配。",
            "答复：感谢专家指出。我们重新区分了“方法面向实时在线应用”和“本文已完成的实验验证”两个层级。题目改为“面向实时在线应用……”，摘要、引言和数据预处理部分均明确说明本文实验为历史监测数据离线回放，未开展现场实时数据流接入和完整在线滚动系统测试。对于原稿中“跨区间迁移”的外推表述，已收敛为“后续多区间数据泛化验证方向”，不再作为本文已验证贡献。同时，我们在3.2节新增实时计算开销估计：模型参数存储约0.40～3.25 MB，单样本前向推理约0.47～3.68 ms，最大输入/输出缓存约45 KB量级，说明在5 s采样周期下，较短H滚动调用具备实时计算可行性。",
            "修改位置：题目、摘要、英文摘要、引言第1段和末段、1.3节数据预处理、3.2节计算性能、3.3节、3.6节、4.1节、4.2节和结论。"
        ),
        (
            "意见2：引言中的参考文献对多变量时间序列预测和Transformer时序预测前沿的支撑不足，归纳仍偏宽泛，未充分凸显本文与已有盾构参数预测研究之间的差异。",
            "答复：已补充和替换相关文献，并强化本文差异。引言新增近年盾构/TBM参数预测中的注意力混合模型、LSTM-Transformer模型和CNN/GRU/Attention模型文献，说明注意力机制已开始应用于盾构姿态、贯入度、推力和刀盘扭矩等任务；同时指出现有工作多聚焦单一目标、单一预测跨度或少量模型对比。本文的差异在于：面向32维关键掘进参数开展多变量联合预测，在统一直接多步架构下系统扫描5×4组(T,H)配置，并同时讨论精度衰减、计算成本和实时滚动调用需求。",
            "修改位置：引言第2段；参考文献新增[25]–[27]。"
        ),
        (
            "意见3：引言中部分参考文献与正文内容关联度不高，例如文献[25]为“双孔隧道中后掘进盾构对地表沉降的影响”，正文将其用于支撑“多头设计、编码器并行计算”等内容，检查全文引用是否相关。",
            "答复：已删除原不相关文献[25]，并将Transformer多头注意力和并行计算的理论依据统一改引Vaswani等提出的Transformer原文[24]。新的[25]替换为与盾构施工参数预测直接相关的文献，并新增[26]、[27]支撑盾构/TBM深度时序预测研究背景。",
            "修改位置：引言第2段；参考文献[24]–[27]。"
        ),
        (
            "意见4：本文整体撰写方式偏说明性，1.4节中花大量篇幅解释多头自注意力。建议压缩一般性原理介绍，突出研究问题、方法构建等关键学术问题。",
            "答复：已压缩1.4节的一般原理铺陈，保留必要原理精华，并突出与本文问题的对应关系。修订后1.4节重点说明“历史窗口编码—当前状态压缩—未来多步直接输出”的结构、单次前向传播延迟确定的实时滚动调用优势，以及同一上下文向量承担多步输出时在大H下形成的结构瓶颈。复杂度部分保留T对注意力计算、H对输出规模和参数量的影响，用以衔接后文精度—效率实验。",
            "修改位置：1.4节注意力编码预测模型、模型结构与超参数、复杂度与结构约束。"
        ),
        (
            "意见5：文章中关于模型隐维度等关键参数前后表述不一致，例如1.4节中隐维度d为4，而实验设置中隐维度为64。",
            "答复：已核对代码。`BASE/AI-REF/Transformer-TS/02-Processing/config.py`中`d_model=64`，`models/transformer_model.py`中`TransformerForecaster`默认`d_model=64`，实际训练脚本也从配置读取该值。原文d=4为笔误，已统一修正为d=64，并补充2层编码器、4头注意力、前馈维度256、Dropout 0.1、学习率1×10⁻³、批大小64和20个epoch等参数。",
            "修改位置：1.4节模型结构与超参数；第2节训练设置。"
        ),
        (
            "意见6：文中存在图号引用不一致的问题，将预测精度指标图误引为图3（计算性能图）。图10中文图题为“土舱土压”，英文图题为“Cutterhead speed”，中英文图名不一致。",
            "答复：已核对并修正全文图号与图题。预测精度指标图统一为图2，计算性能图为图3；3.1节原“如图3所示”已改为“如图2所示”。图10中文图题已由“土舱土压”改为“刀盘转速”，与英文“Cutterhead speed”以及正文中选取的代表性参数一致。",
            "修改位置：3.1节、图10中文图题及对应英文图题。"
        ),
    ]

    for comment, reply, loc in items:
        add_para(doc, comment, bold=True)
        add_para(doc, reply)
        add_para(doc, loc)
        add_para(doc, "")

    add_para(doc, "再次感谢编辑和审稿专家的宝贵意见。")
    doc.save(RESPONSE)


def write_ris_and_md():
    RIS.write_text(
        """TY  - JOUR
AU  - Zhang, Yuxian
AU  - Ren, Xuhua
AU  - Zhang, Jixun
AU  - Zhang, Yuzhe
AU  - Ma, Zichang
TI  - A novel workflow including denoising and hybrid deep learning model for shield tunneling construction parameter prediction
JO  - Engineering Applications of Artificial Intelligence
VL  - 133
SP  - 108103
PY  - 2024
DO  - 10.1016/j.engappai.2024.108103
ER  -

TY  - JOUR
AU  - Dai, Linfabao
AU  - Chen, Wenming
AU  - Xiao, Mingqing
AU  - Sun, Wenhao
AU  - Wang, Zhengzheng
TI  - Prediction of super-large diameter shield attitude based on LSTM-Transformer
JO  - Scientific Reports
VL  - 15
SP  - 15725
PY  - 2025
DO  - 10.1038/s41598-025-98428-8
ER  -

TY  - JOUR
AU  - Yao, Changrui
AU  - Kong, Xiangxun
AU  - Tang, Liang
AU  - Ling, Xianzhang
AU  - Tang, Wenchong
TI  - A multivariate time series prediction model for TBM excavation parameters using a Convolution-GRU-Attention neural network
JO  - Applied Sciences
VL  - 16
IS  - 6
SP  - 2964
PY  - 2026
DO  - 10.3390/app16062964
ER  -
""",
        encoding="utf-8",
    )

    MD.write_text(
        """# 审稿意见修改说明

**稿号**：gl2026-1222  
**论文标题**：面向实时在线应用的盾构掘进关键参数多步预测研究  
**修改稿文件**：`基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx`  
**审稿意见文件**：`gl2026-1222审稿意见.docx`  
**答复文件**：`gl2026-1222审稿意见-答复.docx`  
**修改日期**：2026-06-20  

## 修改说明

1. 针对“实时/在线/滚动/跨区间迁移”表述与实验不匹配的问题，已将全文表述调整为“方法面向实时在线应用，本文采用历史数据离线回放验证核心预测器”。同时新增实时计算开销估计：模型参数存储约0.40～3.25 MB，单样本推理约0.47～3.68 ms，最大输入/输出缓存约45 KB量级；原稿中跨区间迁移外推已收敛为后续多区间数据泛化验证方向。
2. 针对引言支撑不足和研究差异不突出的问题，补充盾构/TBM参数预测中Attention、LSTM-Transformer、CNN/GRU/Attention等相关文献，并强调本文差异在于多变量联合、直接多步、5×4组(T,H)联合扫描及实时计算开销评估。
3. 针对原文献[25]不相关的问题，已删除原不相关文献并以盾构施工参数预测相关文献替换；Transformer多头与并行计算依据统一改引[24]。
4. 针对1.4节说明性过强的问题，已压缩一般性原理介绍，保留“历史窗口编码—当前状态压缩—未来多步直接输出”、复杂度缩放和大H结构瓶颈等与本文科学问题直接相关的原理。
5. 针对隐维度不一致的问题，已核对代码并统一为d=64；相关配置为2层编码器、4头注意力、前馈维度256、Dropout 0.1、学习率1×10⁻³、批大小64、20个epoch。
6. 针对图号和图题问题，已将预测精度图引用修正为图2，将图10中文图题修正为“刀盘转速”，与英文Cutterhead speed一致。

## Zotero 文献导入说明

Cursor 当前无法直接写入本机 Zotero 文件夹；已在 `_revision_work/new_references_for_zotero_import.ris` 生成新增文献的RIS导入文件，可导入到该论文对应的 Zotero 文件夹。
""",
        encoding="utf-8",
    )


def main():
    update_manuscript()
    update_response_doc()
    write_ris_and_md()
    print("Deep review revision complete.")
    print(f"RIS: {RIS}")


if __name__ == "__main__":
    main()
