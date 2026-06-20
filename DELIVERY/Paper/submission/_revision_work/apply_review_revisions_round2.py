# -*- coding: utf-8 -*-
"""Second round of reviewer revisions."""
from docx import Document
from docx.shared import RGBColor

DOC_PATH = (
    r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission"
    r"\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx"
)
RED = RGBColor(255, 0, 0)


def set_run_red(run):
    run.font.color.rgb = RED


def replace_paragraph_text(paragraph, new_text):
    if paragraph.text == new_text:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        set_run_red(paragraph.runs[0])
    else:
        run = paragraph.add_run(new_text)
        set_run_red(run)
    return True


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def main():
    doc = Document(DOC_PATH)
    log = []

    # Compress structure paragraph (keep key params d=64)
    struct_new = (
        "模型结构与超参数：输入 $X\\in\\mathbb{R}^{T\\times F}$ 经线性嵌入（隐维度 $d=64$）"
        "与正弦位置编码后，经 2 层编码器（4 头注意力，前馈维度 256，Dropout 0.1）得到最后一时刻"
        "隐向量 $h_T$，再经线性层映射为 $\\hat{Y}\\in\\mathbb{R}^{H\\times F}$（直接多步）。"
        "所有未来步共享同一 $h_T$，无解码器与自回归；该设计便于在统一架构下比较不同 (T,H)，"
        "也使 H 增大时“单向量承载多步输出”成为主要结构约束（见第 3 节）。"
    )
    for para in doc.paragraphs:
        if para.text.startswith("结构与数据流：输入序列"):
            if replace_paragraph_text(para, struct_new):
                log.append("1.4节：压缩“结构与数据流”段落")
            break

    # Merge complexity + four properties + info theory into one shorter paragraph
    merge_marker = "计算复杂度与 (T, H) 的缩放"
    merge_new = (
        "复杂度与结构约束：自注意力每层复杂度约 $O(T^2 d)$，输出层参数量与 $H$ 线性相关，"
        "故训练/推理时间随 T 约二次、随 H 约一次增长；注意力矩阵显存约 $O(T^2)$。"
        "结合实验现象，可归纳三点与本文直接相关的结构含义：（1）$H=1$ 时单向量解码压力最小，"
        "与 3.1 节高 R² 一致；（2）$H$ 增大时同一 $h_T$ 需同时拟合多步，容量摊薄且远期不确定性上升，"
        "与长 H 配置精度骤降一致；（3）T 继续增大对短 H 的边际收益递减，与算力二次增长形成权衡。"
        "从信息瓶颈视角，定长 $h_T$ 编码多步未来存在容量上限，与可预测性时域衰减共同解释 H 的主导作用。"
    )
    started = False
    for para in doc.paragraphs:
        if para.text.startswith(merge_marker):
            replace_paragraph_text(para, merge_new)
            started = True
            log.append("1.4节：合并复杂度、性质与理论解读段落")
            continue
        if started and (
            para.text.startswith("架构带来的四条性质")
            or para.text.startswith("从信息与表示角度的理论解读")
        ):
            replace_paragraph_text(para, "")
            log.append("1.4节：删除冗余段落（已合并）")

    replacements = [
        (
            "推理时间决定实时或准实时部署时的延迟上限",
            "推理时间决定潜在在线部署场景下的延迟上限（本文未验证在线系统）",
        ),
        (
            "部署上宜以较短H配合滚动递推与实时修正，而非一次输出极长未来",
            "从离线结果推断，潜在在线部署宜以较短 H 配合滚动递推，而非一次输出极长未来（未经验证）",
        ),
        (
            "在线则与滚动更新配合",
            "潜在在线场景可与滚动更新配合（本文未开展在线实验）",
        ),
        (
            "以兼顾精度与实时性",
            "以兼顾精度与低延迟需求",
        ),
    ]
    for para in doc.paragraphs:
        for old, new in replacements:
            if old in para.text:
                replace_in_paragraph(para, old, new)
                log.append(f"替换：{old[:30]}...")

    # Remove uncited reference [25]
    for para in doc.paragraphs:
        if para.text.startswith("[25]"):
            replace_paragraph_text(para, "")
            log.append("参考文献：删除正文未再引用的文献[25]")
            break

    doc.save(DOC_PATH)
    print("Round2 changes:", len(log))
    for x in log:
        print(x)


if __name__ == "__main__":
    main()
