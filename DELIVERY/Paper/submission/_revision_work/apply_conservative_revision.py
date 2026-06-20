# -*- coding: utf-8 -*-
"""Balanced revision: original base + substantive edits with expanded red highlighting."""
from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

BASE = Path(r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission")
ORIG = BASE / "基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿.docx"
REV = BASE / "基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx"

RED = RGBColor(0xFF, 0x00, 0x00)


def para_replace(paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    i = full.index(old)
    before, after = full[:i], full[i + len(old):]
    paragraph.clear()
    if before:
        paragraph.add_run(before)
    run = paragraph.add_run(new)
    run.font.color.rgb = RED
    if after:
        paragraph.add_run(after)
    return True


def para_all_red(paragraph, text: str):
    paragraph.clear()
    r = paragraph.add_run(text)
    r.font.color.rgb = RED


def find_para(doc, substring: str):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def insert_red_after(paragraph, text: str):
    new_el = deepcopy(paragraph._element)
    paragraph._element.addnext(new_el)
    np = Paragraph(new_el, paragraph._parent)
    np.clear()
    r = np.add_run(text)
    r.font.color.rgb = RED


# substring replacements: (marker, old, new) — changed parts red
REPLACEMENTS = [
  # 摘要
  (
    "构建输入长度T与预测跨度H协同配置评估框架",
    "构建输入长度T与预测跨度H协同配置评估框架",
    "构建输入长度 T 与预测跨度 H 协同配置评估框架；实验采用苏州地铁区间历史监测数据离线回放，验证模型结构与 T–H 配置规律，现场实时数据流接入与在线滚动系统测试未纳入本文范围",
  ),
  (
    "为盾构多变量智能预报的在线部署与跨区间迁移提供可复核的方法依据",
    "为盾构多变量智能预报的在线部署与跨区间迁移提供可复核的方法依据",
    "为盾构多变量智能预报的实时在线预测方法配置提供可复核依据；跨区间迁移与现场部署验证留待后续多区间数据与工程化测试",
  ),
  # 引言
  (
    "实时或短时预测可为掘进参数调整与风险预警提供依据[2]",
    "实时或短时预测可为掘进参数调整与风险预警提供依据[2]",
    "在 5 s 采样序列上选取预测步数 H 即可获得秒至分钟量级的向前预报，可为掘进参数调整与风险预警提供依据[2]。本文所称“实时/在线”指模型面向连续入流数据、在新观测到达后滑动窗口快速重算并输出短时前瞻结果的应用形态，而非本文已完成现场部署或实时数据流联调验证",
  ),
  (
    "编码器可并行计算[25]",
    "编码器可并行计算[25]",
    "编码器可并行计算[24]",
  ),
  (
    "在盾构等工业应用中，仍缺乏该类模型应用的案例",
    "在盾构等工业应用中，仍缺乏该类模型应用的案例",
    "在盾构等工业应用中，已有研究将注意力机制用于姿态、贯入度、推力或刀盘扭矩等预测[25–27]，但仍缺乏该类模型在同一区间多变量关键参数上、直接多步架构下的系统对比",
  ),
  (
    "提炼面向部署的参数选型法则、性能边界与可迁移实验范式",
    "提炼面向部署的参数选型法则、性能边界与可迁移实验范式",
    "提炼面向实时在线预测应用的参数选型法则与可复现实验范式；现场在线部署、实时流测试及跨区间泛化验证不属于本文实验范围",
  ),
  # 数据
  (
    "部署时采用滚动预测（窗口随新观测滑动更新。离线实验聚焦",
    "部署时采用滚动预测（窗口随新观测滑动更新。离线实验聚焦",
    "方法设计上可按在线滚动方式使用（窗口随新观测滑动更新）；本文离线实验聚焦",
  ),
  (
    "对不同 (T, H) 的直接多步设置做并行对照",
    "对不同 (T, H) 的直接多步设置做并行对照",
    "对不同 (T, H) 的直接多步设置做并行对照，未接入现场实时数据流，也未对包含采集、缓冲、触发与展示的完整在线系统进行独立测试",
  ),
  # 1.4 隐维度
  (
    "本文d=4",
    "本文d=4",
    "本文 d=64",
  ),
  # 评估指标
  (
    "推理时间决定实时或准实时部署时的延迟上限",
    "推理时间决定实时或准实时部署时的延迟上限",
    "推理时间决定潜在在线部署时的延迟上限（本文未验证在线系统）",
  ),
  # 实验设计
  (
    "为工程中的配置选型提供实证依据",
    "为工程中的配置选型提供实证依据",
    "为面向实时在线应用的配置选型提供离线回放证据",
  ),
  (
    "20 组实验中，R² 分布于",
    "20 组实验中，R² 分布于",
    "基于离线回放测试集评估，20 组实验中，R² 分布于",
  ),
  # 3.1 图号
  (
    "如图 3 所示，四类精度指标在 20 组配置下的表现以柱状图形式给出",
    "如图 3 所示，四类精度指标在 20 组配置下的表现以柱状图形式给出",
    "如图 2 所示，四类精度指标在 20 组配置下的表现以柱状图形式给出",
  ),
  # 3.6 部署表述
  (
    "部署上宜以较短H配合滚动递推与实时修正，而非一次输出极长未来",
    "部署上宜以较短H配合滚动递推与实时修正，而非一次输出极长未来",
    "面向实时在线使用时更宜以较短 H 进行高频滚动调用，而非一次输出极长未来；本文验证的是核心预测器离线回放表现，完整在线系统测试有待后续开展",
  ),
  # 3.6 理论归纳
  (
    "工程上本文可支持的原则限于：在单次直接多步、固定 H 的离线评估框架与本文数据条件下，优先较短 H 与适度 T；在线则与滚动更新配合",
    "工程上本文可支持的原则限于：在单次直接多步、固定 H 的离线评估框架与本文数据条件下，优先较短 H 与适度 T；在线则与滚动更新配合",
    "工程上本文可支持的原则限于：在单次直接多步、固定 H 的离线回放评估框架与本文数据条件下，优先较短 H 与适度 T；该原则可作为实时在线滚动预测器的配置依据，但现场部署与误差反馈机制仍需另设实验",
  ),
  # 图10
  (
    "图10 土舱土压（右上）",
    "图10 土舱土压（右上）",
    "图10 刀盘转速（右上）",
  ),
  # 4.1
  (
    "可参考本文的 20 组设置进行小规模试验",
    "可参考本文的 20 组设置进行小规模试验",
    "可参考本文的 20 组设置进行小规模离线回放试验",
  ),
  (
    "不宜作为首选",
    "不宜作为首选",
    "不宜作为实时滚动预测器的首选配置",
  ),
  (
    "与滚动刷新输入、取较短 H、随新数据重算的在线用法相配合更为合理",
    "与滚动刷新输入、取较短 H、随新数据重算的在线用法相配合更为合理",
    "与滚动刷新输入、取较短 H、随新数据重算的在线用法相配合更为合理；需说明的是，本文未开展上述在线滚动流程的现场验证",
  ),
  # 4.2 局限
  (
    "结论在更多地层与机型上的泛化能力尚待验证",
    "结论在更多地层与机型上的泛化能力尚待验证",
    "结论在更多地层、机型及跨区间数据上的泛化与迁移能力尚待验证；本文实验为单一区间历史数据离线回放，未包含实时流测试",
  ),
  # 结论
  (
    "本文以春光路站—春秋路站区间盾构掘进关键参数为对象",
    "本文以春光路站—春秋路站区间盾构掘进关键参数为对象",
    "本文以春光路站—春秋路站区间盾构掘进关键参数为对象（基于历史监测数据离线回放实验）",
  ),
  # 英文摘要
  (
    "with conclusions bounded by the data and preprocessing conditions in this work.",
    "with conclusions bounded by the data and preprocessing conditions in this work.",
    "with conclusions bounded by the data and preprocessing conditions in this work. Experiments use offline replay of historical monitoring data; field real-time stream integration, online rolling-system tests, and cross-section transfer validation remain for future work.",
  ),
]

# 整段改写（压缩 1.4，回应审稿人第 4 条）— 整段标红
SEC_14_P31 = (
    "本文采用仅含编码器的注意力模型进行多变量多步预测：输入经线性嵌入（隐维度 d=64）"
    "与正弦位置编码后，送入 2 层多头自注意力编码器（4 头，前馈维度 256，Dropout 0.1），"
    "仅取最后一时间步隐向量经单层线性映射直接输出未来 H 步。"
    "该“历史编码—状态压缩—多步直接输出”结构与 Transformer 编码器思想一致[24]，"
    "无解码器、无自回归，单次前向延迟确定，适合连续采样下作为滚动调用的核心预测器；"
    "同一上下文向量需同时承载 H 步输出，在大 H 下存在结构瓶颈，须结合下文 T–H 联合实验确定适用尺度。"
)

SEC_14_P32 = (
    "编码器通过自注意力使任意两时间步信息路径长度为 1，最后一步表示已聚合全长历史；"
    "输出为线性映射，H=1 时解码压力最小，H 增大时同一表示需同时拟合多步，与 3.1 节精度衰减一致。"
    "自注意力每层复杂度约 O(T²d)，注意力矩阵显存约 O(T²)；输出层参数量与 H 线性相关，"
    "故训练/推理时间随 T 约二次、随 H 约一次增长，长 T 配置计算与内存成本显著上升。"
)

SEC_14_P34_35 = (
    "从信息瓶颈与可预测性时域视角，定长隐向量编码多步未来存在容量上限；"
    "条件熵随预测步长增大通常不减，与实验中 H 主导精度衰减、T 边际收益递减"
    "及长 H 配置精度骤降的现象一致，也为后文配置选型提供理论参照。"
)

INSERT_AFTER = [
    # 引言：与 [25–27] 差异对比（新增整段红字）
    (
        "但仍缺乏该类模型在同一区间多变量关键参数上、直接多步架构下的系统对比",
        "与文献[25]侧重去噪与混合深度学习的施工参数预测、文献[26]面向超大直径盾构姿态的 LSTM-Transformer、"
        "文献[27]以 CNN-GRU-Attention 预测 TBM 掘进参数不同，本文在固定仅编码器直接多步架构下，"
        "对同一盾构区间 32 维关键参数开展 5×4 组 (T,H) 系统扫描，联合报告 MSE/MAE/RMSE/R²、"
        "训练时间、推理时间与单样本推理开销，以支持实时滚动调用场景下的 T–H 配置选型。",
    ),
    # 4.1 边界说明
    (
        "机理或统计融合可作为提升可解释性与稳健性的补充路径",
        "需要明确：本文工程建议基于历史数据离线验证所得，现场实时流接入、跨区间数据迁移及完整在线系统（含 I/O 与系统集成）测试有待后续工程化研究。",
    ),
]

NEW_REFS = [
    "[25]\tZHANG Y X, REN X H, ZHANG J X, ZHANG Y Z, MA Z C. A novel workflow including denoising and hybrid deep learning model for shield tunneling construction parameter prediction[J]. Engineering Applications of Artificial Intelligence, 2024, 133: 108103.",
    "[26]\tDAI L F, CHEN W M, XIAO M Q, SUN W H, WANG Z Z. Prediction of super-large diameter shield attitude based on LSTM-Transformer[J]. Scientific Reports, 2025, 15: 15725.",
    "[27]\tYAO C R, KONG X X, TANG L, LING X Z, TANG W C. A multivariate time series prediction model for TBM excavation parameters using a Convolution-GRU-Attention neural network[J]. Applied Sciences, 2026, 16(6): 2964.",
]

COMPUTE_BLOCK = (
    "进一步统计单样本推理开销：20 组配置参数量约 1.04×10^5～8.51×10^5，"
    "模型存储约 0.40～3.25 MB，单样本推理约 0.47～3.68 ms，输入/输出缓存约 45 KB。"
    "上述指标表明，在本文数据规模与 CPU 推理环境下，模型具备作为滚动调用核心预测器的算力潜力；"
    "未计入现场数据采集、传输、缓冲及人机界面等系统集成开销。"
)


def main():
    doc = Document(str(ORIG))
    target = REV

    ok, fail = 0, 0
    for marker, old, new in REPLACEMENTS:
        p = find_para(doc, marker)
        if p is None:
            print("FAIL find:", marker[:50])
            fail += 1
            continue
        if not para_replace(p, old, new):
            print("FAIL replace:", old[:50])
            fail += 1
        else:
            ok += 1

    for marker, text in INSERT_AFTER:
        p = find_para(doc, marker)
        if p:
            insert_red_after(p, text)
            ok += 1
        else:
            print("FAIL insert marker:", marker[:50])
            fail += 1

    # 1.4 压缩：三段整段标红
    p31 = find_para(doc, "多头自注意力编码预测模型在时序预测中的优势在于")
    if p31:
        para_all_red(p31, SEC_14_P31)
    p32 = find_para(doc, "结构与数据流：输入序列")
    if p32:
        para_all_red(p32, SEC_14_P32)
    p34 = find_para(doc, "架构带来的四条性质")
    if p34:
        para_all_red(p34, SEC_14_P34_35)
        p35 = find_para(doc, "从信息与表示角度的理论解读")
        if p35:
            p35.clear()
    # 删除与压缩段重复的“计算复杂度”原段
    p33 = find_para(doc, "计算复杂度与 (T, H) 的缩放")
    if p33:
        p33.clear()

    # 训练设置：补充与代码一致说明（红字）
    p_train = find_para(doc, "关键模型结构及超参数为隐维度 64")
    if p_train:
        insert_red_after(
            p_train,
            "上述隐维度、层数、头数、前馈维度、Dropout、优化器与学习率等超参数已与代码实现（TransformerForecaster 及训练配置）核对一致。",
        )

    # 3.2 实时开销（新增红字段）
    p32sec = find_para(doc, "内存占用自约 484 MB")
    if p32sec and "单样本推理开销" not in p32sec.text:
        insert_red_after(p32sec, COMPUTE_BLOCK)

    # 参考文献 [25]–[27]
    wrong_ref = find_para(doc, "罗雄文,张文广,梁荣柱.双孔隧道")
    if wrong_ref:
        para_all_red(wrong_ref, NEW_REFS[0])
        el = wrong_ref._element
        for extra in NEW_REFS[1:]:
            new_el = deepcopy(el)
            el.addnext(new_el)
            el = new_el
            np = Paragraph(new_el, wrong_ref._parent)
            np.clear()
            r = np.add_run(extra)
            r.font.color.rgb = RED

    try:
        doc.save(str(target))
    except PermissionError:
        target = BASE / "基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿_新.docx"
        print("WARN: 修改稿被占用，将保存到:", target.name)
        doc.save(str(target))
    print(f"Saved: {target}")
    print(f"replacements ok={ok} fail={fail}")


if __name__ == "__main__":
    main()
