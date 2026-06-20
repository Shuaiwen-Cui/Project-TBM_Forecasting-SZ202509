# -*- coding: utf-8 -*-
"""Refine cross-section transfer wording in the revised manuscript and response."""
from docx import Document
from docx.shared import RGBColor

DOC_PATH = (
    r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission"
    r"\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx"
)
MD_PATH = (
    r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission"
    r"\审稿意见修改说明-gl2026-1222.md"
)

RED = RGBColor(255, 0, 0)


def mark_red(run):
    run.font.color.rgb = RED


def set_para(paragraph, text):
    if paragraph.text == text:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        mark_red(paragraph.runs[0])
    else:
        run = paragraph.add_run(text)
        mark_red(run)
    return True


def replace_marker(doc, marker, text):
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            return set_para(paragraph, text)
    raise RuntimeError(f"marker not found: {marker}")


def main():
    doc = Document(DOC_PATH)
    changes = []

    abstract_cn = (
        "摘要：复杂地层盾构施工中，推进压力、刀盘扭矩、土舱土压等多变量参数呈强耦合、快波动特征；"
        "现场连续采样条件下，掘进参数调控与风险预警需要模型能够随新观测快速滚动更新并给出短时前瞻预测。"
        "面向这一实时在线应用需求，本文构建基于多头自注意力机制的多变量直接多步预测方法，并以输入长度 T 与预测跨度 H 的协同配置为核心开展评估。"
        "需要说明的是，本文实验采用苏州某地铁区间历史监测数据的离线回放方式，重点验证模型结构与 T–H 配置规律，"
        "尚未开展现场实时数据流接入和在线滚动系统测试；原稿涉及的跨区间迁移外推已收敛为后续泛化验证方向。"
        "在 20 组受控实验中，结果表明 H 是性能衰减的主导因子：短时预测稳定且 T 的边际收益递减；"
        "当 H 扩展至 120、360 步时，单次直接多步输出显著退化。"
        "结合信息瓶颈与可预测性时域分析，本文提出短跨度优先、滚动更新调用的配置建议，"
        "为盾构掘进关键参数实时在线预测方法的后续工程化验证提供离线实验依据。"
    )
    if set_para(doc.paragraphs[3], abstract_cn):
        changes.append("摘要：将跨区间迁移改为原稿外推表述的后续泛化验证方向")

    abstract_en = (
        "Abstract: In complex shield tunnelling operations, key multivariate parameters "
        "(e.g., propulsion pressure, cutterhead torque, and chamber pressure) are strongly "
        "coupled and fast-varying. Under continuous field sampling, parameter control and "
        "risk warning require forecasts that can be refreshed rapidly with newly observed "
        "data and provide short-horizon look-ahead information. Oriented to this real-time "
        "online forecasting demand, this study develops a multivariate direct multi-step "
        "forecasting method based on multi-head self-attention and evaluates the coupled "
        "configuration of history window length T and forecast horizon H. It should be "
        "noted that the experiments are conducted as offline replay tests using historical "
        "monitoring data from a Suzhou metro shield section; they validate the model "
        "structure and T-H configuration patterns, but do not include field deployment or "
        "real-time stream integration. The cross-section transfer implication in the "
        "original manuscript is therefore restricted here to future generalization "
        "validation. A controlled benchmark over 20 (T,H) settings shows that H is the "
        "dominant factor of performance degradation: short-horizon forecasting remains "
        "robust, while extending T yields limited marginal gains; when H increases to 120 "
        "and 360 steps, direct one-shot multi-step prediction deteriorates markedly. From "
        "information-bottleneck and temporal-predictability perspectives, this study "
        "suggests a configuration strategy of preferring short horizons with rolling "
        "updates, providing offline evidence for subsequent engineering validation of "
        "real-time online shield-tunnelling parameter forecasting."
    )
    if set_para(doc.paragraphs[9], abstract_en):
        changes.append("英文摘要：同步限定 cross-section transfer 为 future validation")

    intro3 = (
        "针对上述缺口，以苏州某地铁盾构区间工程为背景，本文采用注意力编码与直接多步映射相结合的统一建模框架，"
        "以盾构掘进关键参数预测为应用，系统量化 T 与 H 对精度与效率的联合影响，并从信息瓶颈与可预测性时域两个视角给出与下文实验现象相衔接的解释。"
        "数据来源于苏州相城区春光路站—春秋路站区间盾构施工实时监测系统，该区间沿中市路南行，穿越多类地层并下穿兴隆中市桥、复建寿人桥及沿街建构筑物，工程与环境条件具有代表性。"
        "在固定上述架构下，设置 5 种 T 与 4 种 H、共 20 组配置，考察其对预测精度（指标涵盖MSE、MAE、RMSE、R²）与计算效率（训练时间、推理时间、内存）的联合影响，"
        "提炼面向实时在线预测应用的配置选型依据与离线可复现实验范式；现场在线部署与实时流测试不属于本文实验范围，跨区间泛化/迁移能力留待后续多区间数据验证。"
        "全文结构：第 1 节为工程概况、问题定义、数据与模型及评估指标；第 2 节为实验设计；第 3 节为结果与分析（含架构与可预测性时域角度的解读）；第 4 节为工程启示与研究讨论；第5节为结论。"
    )
    if replace_marker(doc, "跨区间迁移验证不属于本文实验范围", intro3):
        changes.append("引言：跨区间迁移改为后续多区间数据验证")

    app2 = (
        "在部署与扩展方面，该注意力框架在长序列下训练与推理成本较高，对延迟敏感的场景可缩短输入长度或采用轻量化结构。"
        "H 为 120、360 步（约 10 min、30 min 向前预报）时，单次直接多步输出的 R² 整体较差；"
        "因此，若构建实时在线预测系统，宜采用“较短 H + 滑动窗口滚动重算”的调用方式，将本文模型作为核心预测器嵌入现场数据流。"
        "本文尚未完成该系统级部署验证，相关的端到端延迟、异常数据鲁棒性、误差反馈均需在后续工程试验中评估；"
        "跨区间泛化/迁移能力也需依托更多区间和机型数据另行验证。"
        "机理或统计融合可作为提升可解释性与稳健性的补充路径，需在统一指标下另行设计对比。"
    )
    if replace_marker(doc, "跨区间迁移能力需在后续工程试验中评估", app2):
        changes.append("工程启示：跨区间能力单列为更多数据验证问题")

    doc.save(DOC_PATH)

    md = """# 审稿意见修改说明

**稿号**：gl2026-1222  
**论文标题**：面向实时在线应用的盾构掘进关键参数多步预测研究  
**修改稿文件**：`基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx`  
**审稿意见文件**：`gl2026-1222审稿意见.docx`  
**修改日期**：2026-06-20  

---

## 修改标注说明

按编辑部要求，修改稿中经修订的文字已用**红色字体**标出。本文档按专家意见逐条说明修改思路、修改位置和理由。

---

## 一、专家审稿意见与修改对照

### 意见 1：实时/在线/滚动/跨区间迁移表述与实验内容不匹配

**审稿意见**：题目、摘要及正文多处强调“实时预测、在线部署、滚动更新、跨区间迁移”，但文中实际开展的是基于单一区间数据的离线直接多步预测实验，未进行实时数据流测试、在线滚动预测验证以及跨区间迁移实验。

**修改原则**：本研究的模型设计服务于盾构现场连续采样下的实时在线预测需求，因此保留“面向实时在线应用”的研究定位；同时严格区分“方法应用目标”和“本文已完成的实验验证”。对于原稿中“为在线部署与跨区间迁移提供依据”“可迁移实验范式”等外推表述，已改为“跨区间泛化/迁移能力需后续多区间数据验证”，不再作为本文已验证贡献。

| 位置 | 修改内容 |
|------|----------|
| 中文标题 | 改为“**面向实时在线应用**的盾构掘进关键参数多步预测研究”，保留应用导向，但不宣称已完成工程实时系统 |
| 英文标题 | 改为 “for Real-Time Online Applications”，与中文题目一致 |
| 中文摘要 | 明确“面向现场连续采样下的实时在线预测需求”，同时说明实验为“历史监测数据离线回放”；原稿涉及的跨区间迁移外推已收敛为后续泛化验证方向 |
| 英文摘要 | 同步补充 real-time online demand、offline replay tests，并将 cross-section transfer 限定为 future generalization validation |
| 引言首段 | 新增对“实时/在线”的定义：指新观测到达后滑动窗口快速重算并输出短时前瞻结果的应用形态，不等同于本文已完成现场部署 |
| 引言末段 | “在线部署与跨区间迁移提供依据/可迁移实验范式”等表述改为“面向实时在线预测应用的配置选型依据与离线可复现实验范式；跨区间泛化/迁移能力留待后续多区间数据验证” |
| 数据与预处理 | 明确方法设计上可按在线滚动方式使用，但本文实验采用历史数据离线回放，未接入现场实时数据流 |
| 结果与讨论 | 将“短 H + 滚动更新”表述为由离线结果支持的实时在线预测器配置方向，同时说明完整在线系统仍需部署或半实物实验验证 |
| 结论与局限 | 明确“面向实时在线预测宜采用较小 H 高频滚动调用”，但本文尚未完成现场在线部署和实时数据流验证；跨区间泛化需后续数据验证 |

---

### 意见 2：引言参考文献对时序预测前沿及盾构预测差异支撑不足

**修改措施**：在引言中补充本文与既有盾构预测研究的差异：已有研究多集中于单一或少量 (T,H)、特定参数与浅层/循环网络，本文侧重多变量联合、直接多步输出及 T–H 联合扫描；同时将应用范围收紧为同一盾构区间数据与统一架构下的系统对比。

---

### 意见 3：部分参考文献与正文关联度不高（如文献 [25]）

**修改措施**：删去原文中用 [25] 支撑多头设计、编码器并行计算的表述，相关 Transformer 架构依据统一改引 [24]；删除未再被正文引用的原参考文献 [25]。

---

### 意见 4：撰写偏说明性，1.4 节原理介绍过多

**修改措施**：压缩 1.4 节的一般性原理介绍，将重点放在本文模型的直接多步结构、低延迟滚动调用潜力、T/H 复杂度缩放和后文实验现象之间的关系；将原“结构与数据流”“复杂度”“架构性质”“信息论解读”整理为更凝练的模型构建与结构约束说明。

---

### 意见 5：隐维度等关键参数前后不一致

**修改措施**：将 1.4 节误写的隐维度 $d=4$ 统一修正为 $d=64$，与实验设置保持一致。

---

### 意见 6：图号引用不一致；图 10 中英文图题不一致

**修改措施**：3.1 节预测精度图引用由“图 3”改为“图 2”；图 10 中文图题由“土舱土压”改为“刀盘转速”，与英文 “Cutterhead speed” 及正文参数选择一致。

---

## 二、编辑审稿意见

**编辑意见**：同意审稿意见，请作者根据专家意见修改。  
**回应**：已按专家意见完成修改，并在稿件中以红色字体标出修改内容。

---

## 三、提交前自查清单

- [x] 已区分“面向实时在线应用”与“本文离线回放验证”
- [x] 已明确未完成现场实时流接入和在线滚动系统测试
- [x] 已将原稿“跨区间迁移”外推收敛为后续多区间数据验证方向
- [x] 已保留方法服务于实时在线预测的研究定位
- [x] 隐维度 $d=64$ 全文一致
- [x] 图 2/图 3 引用与图 10 中英文图题已修正
- [x] 文献 [25] 误引已纠正
- [x] 1.4 节已压缩并强化与实验问题的衔接
- [ ] 作者联系电话需由作者补充
- [ ] 投稿文件名需按编辑部格式最终命名

---

## 四、建议检索核对的关键词

- `面向实时在线应用`
- `离线回放`
- `跨区间泛化/迁移`
- `后续多区间数据验证`
- `滚动调用`
- `d=64`
- `如图 2 所示`
- `刀盘转速`

"""
    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write(md)

    print("Refined cross-section transfer wording.")
    for change in changes:
        print("-", change)


if __name__ == "__main__":
    main()
