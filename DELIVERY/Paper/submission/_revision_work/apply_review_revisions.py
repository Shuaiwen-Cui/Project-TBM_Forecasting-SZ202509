# -*- coding: utf-8 -*-
"""Apply reviewer revisions to manuscript docx with red font for changes."""
from copy import deepcopy
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = (
    r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission"
    r"\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx"
)

RED = RGBColor(255, 0, 0)


def set_run_red(run):
    run.font.color.rgb = RED


def replace_paragraph_text(paragraph, new_text, mark_red=True):
    """Replace entire paragraph text; all runs marked red if mark_red."""
    if paragraph.text == new_text:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        if mark_red:
            set_run_red(paragraph.runs[0])
    else:
        run = paragraph.add_run(new_text)
        if mark_red:
            set_run_red(run)
    return True


def replace_in_paragraph(paragraph, old, new, mark_red=True):
    if old not in paragraph.text:
        return False
    full = paragraph.text.replace(old, new)
    replace_paragraph_text(paragraph, full, mark_red=mark_red)
    return True


def apply_replacements(doc, replacements):
    """replacements: list of (old, new) applied to any paragraph containing old."""
    count = 0
    for para in doc.paragraphs:
        for old, new in replacements:
            if old in para.text:
                replace_in_paragraph(para, old, new)
                count += 1
    return count


def main():
    doc = Document(DOC_PATH)
    changes_log = []

  # --- Title (comment 1: tone down "实时") ---
    title_new = "基于多头自注意力机制的盾构掘进关键参数多步预测研究"
    if replace_paragraph_text(doc.paragraphs[0], title_new):
        changes_log.append("标题：删去“实时”，与离线实验内容一致")

    title_en_new = (
        "Research on Multi-Step Prediction of Key Shield Tunnelling Parameters "
        "Based on Multi-Head Self-Attention Mechanism"
    )
    if replace_paragraph_text(doc.paragraphs[6], title_en_new):
        changes_log.append("英文标题：删去 Real-Time")

    # --- Abstract CN ---
    abstract_cn_new = (
        "摘要：复杂地层盾构施工中，推进压力、刀盘扭矩、土舱土压等多变量参数呈强耦合、快波动特征。"
        "本文以苏州地铁区间连续监测数据为对象，在离线条件下开展基于多头自注意力机制的直接多步预测研究，"
        "构建输入长度 T 与预测跨度 H 的协同配置评估框架，开展 20 组受控对照实验，统一比较精度与计算代价。"
        "结果表明，H 是性能衰减的主导因子：短时预测稳定且 T 边际收益递减；当 H 扩展至 120、360 步时，"
        "单次直接多步输出显著退化。结合信息瓶颈与可预测性时域分析，形成“短跨度优先、滚动更新可增强前瞻能力”"
        "的配置讨论；上述结论限于本文单一区间数据与离线实验设定，在线滚动验证与跨区间迁移尚待后续研究。"
    )
    if replace_paragraph_text(doc.paragraphs[3], abstract_cn_new):
        changes_log.append("中文摘要：收紧实时/在线/跨区间表述，明确离线实验边界")

    keywords_new = (
        "关键词：隧道工程；盾构掘进参数；多头自注意力；时间序列预测；多时间尺度；深度学习"
    )
    if replace_paragraph_text(doc.paragraphs[4], keywords_new):
        changes_log.append("关键词：删去“实时”")

    # --- Abstract EN ---
    abstract_en_new = (
        "Abstract: In complex shield tunnelling operations, key multivariate parameters "
        "(e.g., propulsion pressure, cutterhead torque, and chamber pressure) are strongly "
        "coupled and fast-varying. Using continuous monitoring data from a Suzhou metro "
        "shield section, this study conducts offline direct multi-step forecasting driven "
        "by multi-head self-attention and evaluates the coupled design of history window "
        "length T and forecast horizon H. A controlled benchmark over 20 (T,H) settings "
        "(5×4) under unified training protocols jointly assesses predictive accuracy and "
        "computational cost. Results show that H is the dominant factor of performance "
        "degradation: short-horizon forecasting remains robust, while extending T yields "
        "limited marginal gains; when H increases to 120 and 360 steps, direct one-shot "
        "multi-step prediction deteriorates markedly. From information-bottleneck and "
        "temporal-predictability perspectives, we discuss a configuration strategy of "
        "preferring short horizons with rolling updates for prospective online use; "
        "conclusions are bounded by single-section offline data and are not validated by "
        "real-time stream tests or cross-section transfer experiments in this work."
    )
    if replace_paragraph_text(doc.paragraphs[9], abstract_en_new):
        changes_log.append("英文摘要：与中文摘要对齐，限定结论边界")

    keywords_en_new = (
        "Key words: Tunnel engineering; Shield excavation parameters; Multi-head "
        "Self-attention; Time series forecasting; Multi-time-scale; Deep learning"
    )
    if replace_paragraph_text(doc.paragraphs[10], keywords_en_new):
        changes_log.append("英文关键词：删去 real-time")

    # --- Introduction para 1 ---
    intro1_old = "实时或短时预测可为掘进参数调整与风险预警提供依据"
    intro1_new = "在 5 s 采样序列上选取预测步数 H 即可获得秒至分钟量级的向前预报，可为掘进参数调整与风险预警提供依据"
    for para in doc.paragraphs:
        if intro1_old in para.text:
            replace_in_paragraph(para, intro1_old, intro1_new)
            changes_log.append("引言首段：弱化“实时”表述")
            break

    # --- Introduction para 2: fix ref [25], enhance gap vs shield work ---
    intro2_marker = "编码器可并行计算[25]"
    for para in doc.paragraphs:
        if intro2_marker in para.text:
            old_text = para.text
            new_text = old_text.replace(
                "自注意力在复杂度 下直接建模任意位置对，无需递归即可实现长程依赖[24]；"
                "多头设计便于多变量与多时间子空间表示，编码器可并行计算[25]。",
                "自注意力在 $O(T^2)$ 复杂度下直接建模任意位置对，无需递归即可实现长程依赖[24]；"
                "多头设计与编码器并行计算亦源于 Transformer 架构[24]。",
            )
            new_text = new_text.replace(
                "仍缺乏该类模型应用的案例以及该类架构下对多种 (T, H) 的系统对比",
                "仍缺乏在同一盾构区间数据与统一架构下对多种 (T, H) 的系统对比",
            )
            new_text = new_text.replace(
                "及从信息瓶颈、可预测性时域等理论对选型规律的解读。",
                "及从信息瓶颈、可预测性时域等理论对选型规律的解读；"
                "与既有盾构掘进参数预测研究（多集中于单一或少量 (T,H)、特定参数与浅层/循环网络）相比，"
                "本文侧重多变量联合、直接多步与 T–H 联合扫描。",
            )
            if new_text != old_text:
                replace_paragraph_text(para, new_text)
                changes_log.append("引言：修正文献[25]误引；凸显与既有盾构预测研究的差异")
            break

    # --- Introduction para 3: deployment wording ---
    intro3_old = "提炼面向部署的参数选型法则、性能边界与可迁移实验范式"
    intro3_new = "提炼面向潜在部署讨论的配置选型法则、性能边界与可复现实验范式（跨区间迁移未在本文验证）"
    for para in doc.paragraphs:
        if intro3_old in para.text:
            replace_in_paragraph(para, intro3_old, intro3_new)
            changes_log.append("引言末段：收紧部署与跨区间表述")
            break

    # --- Data preprocessing: clarify offline vs rolling ---
    data_old = "部署时采用滚动预测（窗口随新观测滑动更新。离线实验聚焦"
    data_new = "在线场景可采用滚动预测（窗口随新观测滑动更新），但本文未开展实时数据流或在线滚动验证；离线实验聚焦"
    for para in doc.paragraphs:
        if data_old in para.text:
            replace_in_paragraph(para, data_old, data_new)
            changes_log.append("数据预处理：区分在线设想与本文离线实验")
            break

    # --- Section 1.4: fix d=4 -> d=64 and compress ---
    for para in doc.paragraphs:
        if "本文d=4" in para.text or "本文 d=4" in para.text:
            replace_in_paragraph(para, "本文d=4", "本文 d=64")
            replace_in_paragraph(para, "本文 d=4", "本文 d=64")
            changes_log.append("1.4节：隐维度 d 由 4 更正为 64")
            break

    # Compress attention model opening paragraph
    attn_para_old_start = "多头自注意力编码预测模型在时序预测中的优势在于"
    attn_para_new = (
        "本文采用仅含编码器的注意力模型进行多变量多步预测：输入经线性嵌入与位置编码后，"
        "经多层多头自注意力编码器得到最后一时刻隐向量，再经单层线性映射直接输出未来 H 步（直接多步）。"
        "该设计与 Vaswani 等提出的 Transformer 编码器思想一致[24]，在本文中用于在统一架构下公平比较不同 (T, H)。"
        "自注意力可在 $O(T^2)$ 下建模任意位置依赖并支持并行计算；取最后时间步表示再映射到 $H \\times F$ 输出，"
        "使 T、H 对精度与算力的缩放规律可直接进入后文实验分析（详见 1.4 后段及第 3 节）。"
    )
    for para in doc.paragraphs:
        if attn_para_old_start in para.text:
            replace_paragraph_text(para, attn_para_new)
            changes_log.append("1.4节：压缩一般性原理介绍，突出本文模型与实验衔接")
            break

    # --- Figure reference fix: accuracy should be Fig 2 not Fig 3 ---
    for para in doc.paragraphs:
        if "如图 3 所示，四类精度指标" in para.text:
            replace_in_paragraph(para, "如图 3 所示，四类精度指标", "如图 2 所示，四类精度指标")
            changes_log.append("3.1节：精度图引用由图3改为图2")
            break

    # --- Fig 10 caption fix ---
    for para in doc.paragraphs:
        if "图10 土舱土压" in para.text:
            replace_in_paragraph(
                para,
                "图10 土舱土压（右上）：真实值与预测值（预测长度 360 步）",
                "图10 刀盘转速：真实值与预测值（预测长度 360 步）",
            )
            changes_log.append("图10：中文图题与英文及正文参数一致（刀盘转速）")
            break

    # --- Section 3.3 rolling wording ---
    rolling_old = "结果显示“短预测长度+ 滚动更新”相比于一次性长预测效果更佳。"
    rolling_new = (
        "结果支持在潜在在线场景中采用较短 H 并配合滚动更新，而非单次直接输出极大 H；"
        "该判断来自离线直接多步对照，非本文已完成的在线验证。"
    )
    for para in doc.paragraphs:
        if rolling_old in para.text:
            replace_in_paragraph(para, rolling_old, rolling_new)
            changes_log.append("3.3节：滚动更新表述改为讨论性、非已验证结论")
            break

    # --- Section 4.1 deployment ---
    deploy_old = "实时或轻量部署可优先"
    deploy_new = "若面向低延迟部署讨论，可优先"
    for para in doc.paragraphs:
        if deploy_old in para.text:
            replace_in_paragraph(para, deploy_old, deploy_new)
            changes_log.append("4.1节：部署建议改为条件性讨论")
            break

    deploy2_old = "与滚动刷新输入、取较短 H、随新数据重算的在线用法相配合更为合理"
    deploy2_new = (
        "从离线结果推断，若用于在线场景，宜与滚动刷新输入、取较短 H、随新数据重算相结合；"
        "本文未开展相应在线实验"
    )
    for para in doc.paragraphs:
        if deploy2_old in para.text:
            replace_in_paragraph(para, deploy2_old, deploy2_new)
            changes_log.append("4.1节：在线用法标明为推断而非验证")
            break

    # --- Limitations section ---
    limit_old = "结论在更多地层与机型上的泛化能力尚待验证"
    limit_new = "结论在更多地层、机型及跨区间数据上的泛化与在线滚动验证尚待研究"
    for para in doc.paragraphs:
        if limit_old in para.text:
            replace_in_paragraph(para, limit_old, limit_new)
            changes_log.append("局限：补充跨区间与在线验证")
            break

    limit2_old = "未来研究应结合更多区间与机型数据验证泛化能力"
    limit2_new = "未来研究应结合更多区间与机型数据验证泛化能力，并开展实时数据流与跨区间迁移实验"
    for para in doc.paragraphs:
        if limit2_old in para.text:
            replace_in_paragraph(para, limit2_old, limit2_new)
            changes_log.append("未来工作：明确在线与跨区间实验")
            break

    # --- Conclusion item 2 ---
    concl_old = "与在线取较小 H、滚动修正的部署思路一致"
    concl_new = "与潜在在线场景中取较小 H、滚动修正的思路相一致（本文未验证在线部署）"
    for para in doc.paragraphs:
        if concl_old in para.text:
            replace_in_paragraph(para, concl_old, concl_new)
            changes_log.append("结论：限定在线部署表述")
            break

    concl_rt_old = "实时或资源受限场景宜采用较短 T 与较小 H"
    concl_rt_new = "资源受限场景宜采用较短 T 与较小 H"
    for para in doc.paragraphs:
        if concl_rt_old in para.text:
            replace_in_paragraph(para, concl_rt_old, concl_rt_new)
            changes_log.append("结论：删去未验证的“实时”场景断言")
            break

    doc.save(DOC_PATH)
    print("Saved:", DOC_PATH)
    print("Changes applied:", len(changes_log))
    for i, c in enumerate(changes_log, 1):
        print(f"  {i}. {c}")

    # Write log for md generation
    with open(
        DOC_PATH.replace(".docx", "_changes_log.txt"),
        "w",
        encoding="utf-8",
    ) as f:
        f.write("\n".join(changes_log))


if __name__ == "__main__":
    main()
