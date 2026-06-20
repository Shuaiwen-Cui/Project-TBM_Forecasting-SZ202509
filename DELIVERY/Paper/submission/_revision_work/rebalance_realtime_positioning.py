# -*- coding: utf-8 -*-
"""Rebalance realtime/online positioning in the revised manuscript."""
from docx import Document
from docx.shared import RGBColor

DOC_PATH = (
    r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission"
    r"\基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx"
)
MD_PATH = (
    r"E:\PROJ\Project-TBM_Forecasting-SZ202509\DELIVERY\Paper\submission"
    r"\审稿意见修改说明-gl2026-1222.md"
)

RED = RGBColor(255, 0, 0)


def red(run):
    run.font.color.rgb = RED


def set_para_text(paragraph, text):
    if paragraph.text == text:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        red(paragraph.runs[0])
    else:
        run = paragraph.add_run(text)
        red(run)
    return True


def replace_containing(doc, marker, text):
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            return set_para_text(paragraph, text)
    raise RuntimeError(f"marker not found: {marker}")


def replace_containing_any(doc, markers, text):
    for paragraph in doc.paragraphs:
        if any(marker in paragraph.text for marker in markers):
            return set_para_text(paragraph, text)
    raise RuntimeError(f"markers not found: {markers}")


def main():
    doc = Document(DOC_PATH)
    changes = []

    if set_para_text(doc.paragraphs[0], "面向实时在线应用的盾构掘进关键参数多步预测研究"):
        changes.append("题目：改为“面向实时在线应用”，保留研究定位但避免宣称已完成现场实时验证")

    if set_para_text(
        doc.paragraphs[6],
        "Multi-Step Prediction of Key Shield Tunnelling Parameters for Real-Time Online Applications Based on Multi-Head Self-Attention",
    ):
        changes.append("英文题目：同步体现 real-time online applications 的应用导向")

    abstract_cn = (
        "摘要：复杂地层盾构施工中，推进压力、刀盘扭矩、土舱土压等多变量参数呈强耦合、快波动特征；"
        "现场连续采样条件下，掘进参数调控与风险预警需要模型能够随新观测快速滚动更新并给出短时前瞻预测。"
        "面向这一实时在线应用需求，本文构建基于多头自注意力机制的多变量直接多步预测方法，并以输入长度 T 与预测跨度 H 的协同配置为核心开展评估。"
        "需要说明的是，本文实验采用苏州某地铁区间历史监测数据的离线回放方式，重点验证模型结构与 T–H 配置规律，"
        "尚未开展现场实时数据流接入、在线滚动系统测试或跨区间迁移实验。"
        "在 20 组受控实验中，结果表明 H 是性能衰减的主导因子：短时预测稳定且 T 的边际收益递减；"
        "当 H 扩展至 120、360 步时，单次直接多步输出显著退化。"
        "结合信息瓶颈与可预测性时域分析，本文提出短跨度优先、滚动更新调用的配置建议，"
        "为盾构掘进关键参数实时在线预测方法的后续工程化验证提供离线实验依据。"
    )
    if set_para_text(doc.paragraphs[3], abstract_cn):
        changes.append("中文摘要：重写为“目标面向实时在线，证据来自离线回放”的平衡表述")

    if set_para_text(
        doc.paragraphs[4],
        "关键词：隧道工程；盾构掘进参数；多头自注意力；时间序列实时预测；多时间尺度；深度学习",
    ):
        changes.append("关键词：恢复“时间序列实时预测”的方法定位")

    abstract_en = (
        "Abstract: In complex shield tunnelling operations, key multivariate parameters "
        "(e.g., propulsion pressure, cutterhead torque, and chamber pressure) are strongly "
        "coupled and fast-varying. Under continuous field sampling, parameter control and "
        "risk warning require forecasts that can be refreshed rapidly with newly observed "
        "data and provide short-horizon look-ahead information. Oriented to this real-time "
        "online forecasting demand, this study develops a multivariate direct multi-step "
        "forecasting method based on multi-head self-attention and evaluates the coupled "
        "configuration of history window length T and forecast horizon H. It should be "
        "noted that the experiments are conducted as offline replay tests using historical "
        "monitoring data from a Suzhou metro shield section; they validate the model "
        "structure and T-H configuration patterns, but do not include field deployment, "
        "real-time stream integration, online rolling-system tests, or cross-section "
        "transfer experiments. A controlled benchmark over 20 (T,H) settings shows that H "
        "is the dominant factor of performance degradation: short-horizon forecasting "
        "remains robust, while extending T yields limited marginal gains; when H increases "
        "to 120 and 360 steps, direct one-shot multi-step prediction deteriorates markedly. "
        "From information-bottleneck and temporal-predictability perspectives, this study "
        "suggests a configuration strategy of preferring short horizons with rolling "
        "updates, providing offline evidence for subsequent engineering validation of "
        "real-time online shield-tunnelling parameter forecasting."
    )
    if set_para_text(doc.paragraphs[9], abstract_en):
        changes.append("英文摘要：同步补充 real-time online demand 与 offline replay 边界")

    if set_para_text(
        doc.paragraphs[10],
        "Key words: Tunnel engineering; Shield excavation parameters; Multi-head Self-attention; Real-time time series forecasting; Multi-time-scale; Deep learning",
    ):
        changes.append("英文关键词：恢复 real-time time series forecasting")

    intro1 = (
        "盾构在复杂地层中掘进时，贯入度、推力、刀盘扭矩等关键参数的准确预测对施工安全与效率至关重要[1]；"
        "尤其在软硬交替、多类地层交汇或下穿敏感建构筑物等条件下，在 5 s 采样序列上选取预测步数 H 即可获得秒至分钟量级的向前预报，"
        "可为掘进参数调整与风险预警提供依据[2]。本文所称“实时/在线”主要指模型面向连续入流数据，"
        "能够在新观测到达后以滑动窗口快速重算并输出短时前瞻结果的应用形态，而非本文已经完成现场部署或实时数据流联调。"
        "盾构推进、刀盘、土压等多子系统强耦合，多变量时序既呈现强耦合又呈现长程依赖，是数据驱动多变量时序预测的典型应用场景[3,4]。"
        "工业与土木等领域中，“用多长历史（输入长度 T）、预测多远（预测长度 H）”均制约预测精度与计算成本，"
        "但在同一模型与数据下系统量化 T、H 对精度与效率的联合影响、并从可预测性理论角度加以解释的工作仍较缺乏[5]。"
    )
    if replace_containing_any(
        doc,
        ["本文所称“实时/在线”", "在 5 s 采样序列上选取预测步数 H"],
        intro1,
    ):
        changes.append("引言首段：新增“实时/在线”的定义边界")

    intro3 = (
        "针对上述缺口，以苏州某地铁盾构区间工程为背景，本文采用注意力编码与直接多步映射相结合的统一建模框架，"
        "以盾构掘进关键参数预测为应用，系统量化 T 与 H 对精度与效率的联合影响，并从信息瓶颈与可预测性时域两个视角给出与下文实验现象相衔接的解释。"
        "数据来源于苏州相城区春光路站—春秋路站区间盾构施工实时监测系统，该区间沿中市路南行，穿越多类地层并下穿兴隆中市桥、复建寿人桥及沿街建构筑物，工程与环境条件具有代表性。"
        "在固定上述架构下，设置 5 种 T 与 4 种 H、共 20 组配置，考察其对预测精度（指标涵盖MSE、MAE、RMSE、R²）与计算效率（训练时间、推理时间、内存）的联合影响，"
        "提炼面向实时在线预测应用的配置选型依据与离线可复现实验范式；现场在线部署、实时流测试和跨区间迁移验证不属于本文实验范围。"
        "全文结构：第 1 节为工程概况、问题定义、数据与模型及评估指标；第 2 节为实验设计；第 3 节为结果与分析（含架构与可预测性时域角度的解读）；第 4 节为工程启示与研究讨论；第5节为结论。"
    )
    if replace_containing(doc, "提炼面向潜在部署讨论", intro3):
        changes.append("引言末段：将“潜在部署讨论”改为“面向实时在线应用的离线证据”")

    data_para = (
        "数据集预处理：数据来源于该区间盾构施工实时监测系统，采样间隔 5 s。"
        "该间隔在工程上对应盾构推进与刀盘控制的典型操作节奏，既能捕捉地层与机械交互在数十秒尺度上的变化，又避免过高采样带来的冗余与存储负担；"
        "在方法上则直接决定了每一步所对应的物理时间，从而将抽象的 T、H（步数）与“用多长历史、预测多远”（分钟/秒）建立对应关系。"
        "经合并、去缺与格式统一后，得到以时间顺序排列的连续序列；前 5 列为元数据（管片号、行程、日期、时刻、掘进状态），"
        "其后为 32 维数值型掘进参数，包括贯入度、四腔室推进压力、土舱土压、推进油缸总推力、千斤顶速度与行程、刀盘转速与扭矩及 No.1～No.10 刀盘电机扭矩等。"
        "盾构虽有停机拼装阶段，但监测数据仍以 5 s 间隔持续入流，因此方法设计上可按在线滚动方式使用：新观测到达后滑动输入窗口并重新输出未来 H 步预测。"
        "本文实验阶段采用历史数据离线回放来评估单次前向输出跨度 H 与输入长度 T 的影响，未接入现场实时数据流，也未对完整在线滚动系统进行独立测试。"
    )
    if replace_containing(doc, "盾构虽有停机拼装阶段", data_para):
        changes.append("数据预处理：明确方法支持滚动使用，但实验为历史数据离线回放")

    model_para = (
        "本文采用仅含编码器的注意力模型进行多变量多步预测：输入经线性嵌入与位置编码后，经多层多头自注意力编码器得到最后一时刻隐向量，"
        "再经单层线性映射直接输出未来 H 步（直接多步）。该设计与 Vaswani 等提出的 Transformer 编码器思想一致[24]，"
        "在本文中用于在统一架构下公平比较不同 (T, H)。由于预测过程不依赖递归生成，单次前向传播延迟确定，"
        "适合在连续采样场景下作为滚动调用的核心预测器；但本文验证的是该核心预测器在离线回放数据上的精度—成本特征。"
        "自注意力可在 $O(T^2)$ 下建模任意位置依赖并支持并行计算；取最后时间步表示再映射到 $H \\times F$ 输出，"
        "使 T、H 对精度与算力的缩放规律可直接进入后文实验分析（详见 1.4 后段及第 3 节）。"
    )
    if replace_containing(doc, "本文采用仅含编码器的注意力模型", model_para):
        changes.append("模型节：补充低延迟滚动调用的设计动机与验证边界")

    goal_para = (
        "实验目标：在统一数据与模型前提下，系统改变输入序列长度 T 与预测长度 H，量化二者对预测精度（MSE、MAE、RMSE、R²）与计算效率（训练时间、推理时间、内存占用）的影响，"
        "为面向实时在线应用的配置选型提供离线回放证据，并支撑后续对精度–效率权衡、模型适用边界及“信息瓶颈–可预测性时域”解释的讨论。"
        "设计上需满足：（1）T、H 均有足够跨度，以覆盖从“极短历史/单步预测”到“长历史/长预测”的典型场景；"
        "（2）组间除 (T, H) 外其余条件一致，使差异可归因于时间尺度配置；（3）指标同时包含精度与计算成本，以便提炼选型原则。"
    )
    if replace_containing(doc, "为工程中的配置选型提供实证依据", goal_para):
        changes.append("实验目标：说明离线回放证据服务于实时在线配置选型")

    joint_para = (
        "如图 4 与图 5 所示，分别以三维曲面与热力图展示 R² 随输入序列长度与预测长度的变化。"
        "从曲面与热力图可归纳出以下规律：高 R² 区域集中在预测长度 1 步与 6 步两行，且沿输入长度方向（列方向）变化不大，说明在短预测下增加历史长度对精度提升有限；"
        "预测长度 120 步时仅输入 120 步一格为正（0.104），其余输入长度下均为负；预测长度 360 步时整行为负，右下角 (360, 360) 最低。"
        "整体表明预测长度是主导精度的首要因素，输入长度在短预测下为次要因素。"
        "工程选型时可先根据应用需求确定预测长度，再在满足精度前提下尽量缩短输入长度以控制计算成本。"
        "在 H=360 下“T=H”并未改善精度。该离线结果支持将模型用于实时在线场景时优先采用较短 H 并随新观测滚动调用，"
        "而不宜依赖单次直接输出极大 H；但完整在线系统的延迟、稳定性与误差反馈仍需另设部署或半实物实验验证。"
    )
    if replace_containing(doc, "结果支持在潜在在线场景中采用较短 H", joint_para):
        changes.append("3.3节：把滚动更新从“潜在”改为方法使用方向，同时保留验证边界")

    arch_long = (
        "预测长度 120 步时，仅当输入长度同为 120 步时取得正 R²（0.104）。相对其它 T 略好，但绝对精度仍低。"
        "从结构上看，输入 120 步、输出 120 步在步数跨度上对称，编码器与输出层在“跨度”上一致，模型可能更易学习“用 120 步历史概括并预测 120 步未来”的映射；"
        "而输入 6 步、输出 120 步则要求从极短上下文预测长未来，结构上不匹配，精度显著恶化。"
        "预测长度 360 步时即便输入 360 步仍全部为负 R²，且 (360,360) 为最差配置之一，说明“对称跨度”并不保证优于其它 (T, H)。"
        "在本文数据与模型规模下，单向量表示难以支撑 360 步的单次直接多步输出；因此，面向实时在线使用时更合理的方式是以较短 H 进行高频滚动调用，而非一次输出极长未来。"
        "需要强调的是，本文验证的是核心预测器的离线回放表现，尚未把该滚动调用方式作为完整在线系统进行独立测试。"
    )
    if replace_containing(doc, "从离线结果推断，潜在在线部署宜以较短 H", arch_long):
        changes.append("3.6节：重写长H讨论，体现实时在线方法动机与未部署边界")

    theory_para = (
        "从时序预测理论角度，可将上述现象归纳为两点。其一，预测长度主导精度对应“可预测性随预测时域衰减”："
        "在多数实际系统中，条件分布的方差随 h 增大而增大，或有效信号随h衰减，R² 随 H 增大而下降是预期现象；"
        "本模型采用直接多步且无自回归修正，所有未来步共享同一表示，故 H 的负面影响被进一步放大。"
        "其二，上述同量级启发可部分解释 H =120 时 T 过小会显著变差，但不足以解释 H =360 时 (360,360) 的劣势，亦不能替代实证。"
        "工程上本文可支持的原则限于：在单次直接多步、固定 H 的离线回放评估框架与本文数据条件下，优先较短 H 与适度 T；"
        "该原则可作为实时在线滚动预测器的配置依据，但现场部署、实时流测试及误差反馈机制仍需另设实验。"
        "与机理/统计融合或不同滚动实现作严格对照，也需作为后续工作，不宜由本文直接外推。"
    )
    if replace_containing(doc, "潜在在线场景可与滚动更新配合", theory_para):
        changes.append("理论讨论：把离线原则与实时在线配置依据衔接起来")

    app1 = (
        "基于本文实验结果与架构分析，对工程应用提出以下建议：配置选型方面，在保持采样间隔与特征维度一致的前提下，"
        "可参考本文的 20 组设置进行小规模离线回放试验，根据精度与训练/推理时间权衡选择配置。"
        "若面向实时在线预测器的低延迟调用，可优先输入 6 步或 30 步、预测 1 步，在损失较小精度的前提下显著降低计算成本；"
        "精度优先且可接受较长时间训练时可选用输入 120 步、预测 1 步；短期 6 步预测可考虑输入 6 步或 120 步、预测 6 步。"
        "长序列与长预测组合（如 360–120、360–360）计算成本高而精度差，不宜作为实时滚动预测器的首选配置。"
    )
    if replace_containing(doc, "若面向低延迟部署讨论", app1):
        changes.append("4.1节：明确建议对象是实时在线预测器配置，而非已部署系统")

    app2 = (
        "在部署与扩展方面，该注意力框架在长序列下训练与推理成本较高，对延迟敏感的场景可缩短输入长度或采用轻量化结构。"
        "H 为 120、360 步（约 10 min、30 min 向前预报）时，单次直接多步输出的 R² 整体较差；"
        "因此，若构建实时在线预测系统，宜采用“较短 H + 滑动窗口滚动重算”的调用方式，将本文模型作为核心预测器嵌入现场数据流。"
        "本文尚未完成该系统级部署验证，相关的端到端延迟、异常数据鲁棒性、误差反馈和跨区间迁移能力需在后续工程试验中评估。"
        "机理或统计融合可作为提升可解释性与稳健性的补充路径，需在统一指标下另行设计对比。"
    )
    if replace_containing(doc, "从离线结果推断，若用于在线场景", app2):
        changes.append("4.1节：提出合理工程化路径，明确未做系统级部署验证")

    limitation = (
        "本研究存在以下局限：实验限于单一区间、单一机型及固定超参数，结论在更多地层、机型及跨区间数据上的泛化与在线滚动验证尚待研究；"
        "在 H 为 120、360 步（约 10 min、30 min 向前预报）时直接多步预测效果有限。"
        "掘进与停掘片段均包含在训练与测试序列中，典型曲线显示模型对两类片段均能跟踪；若需单独评价纯掘进段表现，可另设状态筛选作对照，留待后续。"
        "后续将在更多地层与机型上验证泛化能力，并围绕实时在线应用进一步开展数据流接入、滚动预测系统测试、轻量化注意力编码、状态条件化数据构建及与其他模型融合等研究，以兼顾精度、低延迟与工程鲁棒性。"
    )
    if replace_containing(doc, "以兼顾精度与低延迟需求", limitation):
        changes.append("局限：把后续工作具体化为实时在线系统验证")

    conclusion2 = (
        "（2）在 H=120 步时，仅 T=120 步时 R² 勉强为正，其余为负；在 H=360 步时全部配置为负，且 T=H=360 并非改善而为全局最劣之一。"
        "单次大 H 直接多步在本设定下效果差，说明面向实时在线预测时宜采用较小 H 的高频滚动调用，而不是一次性输出过长未来；"
        "但本文尚未完成现场在线部署和实时数据流验证。与机理或统计模型融合等可作为后续受控对比内容，非本文已验证结论。"
        "计算成本随序列与预测长度显著增加，资源受限或低延迟场景宜采用较短 T 与较小 H，离线高精度分析可选用输入 120 步、预测 1 步的配置。"
    )
    if replace_containing(doc, "与潜在在线场景中取较小 H", conclusion2):
        changes.append("结论：保留实时在线方法建议，并明确未做现场在线验证")

    doc.save(DOC_PATH)

    md = """# 审稿意见修改说明

**稿号**：gl2026-1222  
**论文标题**：面向实时在线应用的盾构掘进关键参数多步预测研究  
**修改稿文件**：`基于 Transformer 的盾构掘进关键参数多时间尺度实时预测研究-投稿-修改稿.docx`  
**审稿意见文件**：`gl2026-1222审稿意见.docx`  
**修改日期**：2026-06-20  

---

## 修改标注说明

按编辑部要求，修改稿中经修订的文字已用**红色字体**标出。本文档按专家意见逐条说明修改思路、修改位置和理由。

---

## 一、专家审稿意见与修改对照

### 意见 1：实时/在线/滚动/跨区间表述与实验内容不匹配

**审稿意见**：题目、摘要及正文多处强调“实时预测、在线部署、滚动更新、跨区间迁移”，但文中实际开展的是基于单一区间数据的离线直接多步预测实验，未进行实时数据流测试、在线滚动预测验证以及跨区间迁移实验。

**修改原则**：本研究的模型设计确实服务于盾构现场连续采样下的实时在线预测需求，因此不宜简单删除“实时/在线”这一研究定位；但需把“方法面向实时在线应用”与“本文已完成现场部署验证”严格区分。修订后全文统一采用“面向实时在线应用、离线回放验证核心预测器”的表述框架。

| 位置 | 修改内容 |
|------|----------|
| 中文标题 | 改为“**面向实时在线应用**的盾构掘进关键参数多步预测研究”，保留应用导向，但不表述为已完成工程实时系统 |
| 英文标题 | 改为 “for Real-Time Online Applications”，与中文题目一致 |
| 中文摘要 | 明确“面向现场连续采样下的实时在线预测需求”，同时说明实验为“**历史监测数据离线回放**”，未开展现场实时流接入、在线滚动系统测试或跨区间迁移 |
| 英文摘要 | 同步补充 real-time online demand 与 offline replay tests 的边界 |
| 关键词 | 恢复“时间序列实时预测”，因其是方法应用定位；但摘要和正文均明确未做现场部署验证 |
| 引言首段 | 新增对“实时/在线”的定义：指新观测到达后滑动窗口快速重算并输出短时前瞻结果的应用形态，不等同于本文已完成现场部署 |
| 数据与预处理 | 明确方法设计上可按在线滚动方式使用，但本文实验采用历史数据离线回放，未接入现场实时数据流 |
| 模型方法 | 补充“单次前向传播延迟确定，适合连续采样场景下滚动调用”，并说明本文验证的是核心预测器的离线精度—成本特征 |
| 结果与讨论 | 将“短 H + 滚动更新”表述为由离线结果支持的实时在线预测器配置方向，同时说明完整在线系统仍需部署或半实物实验验证 |
| 结论与局限 | 明确“面向实时在线预测宜采用较小 H 高频滚动调用”，但本文尚未完成现场在线部署和实时数据流验证 |

---

### 意见 2：引言参考文献对时序预测前沿及盾构预测差异支撑不足

**修改措施**：在引言中补充本文与既有盾构预测研究的差异：已有研究多集中于单一或少量 (T,H)、特定参数与浅层/循环网络，本文侧重多变量联合、直接多步输出及 T–H 联合扫描；同时将应用范围收紧为同一盾构区间数据与统一架构下的系统对比。

---

### 意见 3：部分参考文献与正文关联度不高（如文献 [25]）

**修改措施**：删去原文中用 [25] 支撑多头设计、编码器并行计算的表述，相关 Transformer 架构依据统一改引 [24]；删除未再被正文引用的原参考文献 [25]。

---

### 意见 4：撰写偏说明性，1.4 节原理介绍过多

**修改措施**：压缩 1.4 节的一般性原理介绍，将重点放在本文模型的直接多步结构、低延迟滚动调用潜力、T/H 复杂度缩放和后文实验现象之间的关系；将原“结构与数据流”“复杂度”“架构性质”“信息论解读”整理为更凝练的模型构建与结构约束说明。

---

### 意见 5：隐维度等关键参数前后不一致

**修改措施**：将 1.4 节误写的隐维度 $d=4$ 统一修正为 $d=64$，与实验设置保持一致。

---

### 意见 6：图号引用不一致；图 10 中英文图题不一致

**修改措施**：3.1 节预测精度图引用由“图 3”改为“图 2”；图 10 中文图题由“土舱土压”改为“刀盘转速”，与英文 “Cutterhead speed” 及正文参数选择一致。

---

## 二、编辑审稿意见

**编辑意见**：同意审稿意见，请作者根据专家意见修改。  
**回应**：已按专家意见完成修改，并在稿件中以红色字体标出修改内容。

---

## 三、提交前自查清单

- [x] 已区分“面向实时在线应用”与“本文离线回放验证”
- [x] 已明确未完成现场实时流接入、在线滚动系统测试和跨区间迁移
- [x] 已保留方法服务于实时在线预测的研究定位
- [x] 隐维度 $d=64$ 全文一致
- [x] 图 2/图 3 引用与图 10 中英文图题已修正
- [x] 文献 [25] 误引已纠正
- [x] 1.4 节已压缩并强化与实验问题的衔接
- [ ] 作者联系电话需由作者补充
- [ ] 投稿文件名需按编辑部格式最终命名

---

## 四、建议检索核对的关键词

- `面向实时在线应用`
- `离线回放`
- `未开展现场实时数据流`
- `滚动调用`
- `d=64`
- `如图 2 所示`
- `刀盘转速`

"""
    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write(md)

    print("Rebalanced realtime positioning.")
    print("Changes:", len(changes))
    for item in changes:
        print("-", item)


if __name__ == "__main__":
    main()
